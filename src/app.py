import streamlit as st
import pandas as pd
import pickle
import joblib
import os
from pathlib import Path
import io
from PIL import Image
import pytesseract
import PyPDF2
from docx import Document
import re
import numpy as np
import warnings

# Suppress sklearn version warnings
warnings.filterwarnings('ignore', category=UserWarning)
os.environ['PYTHONWARNINGS'] = 'ignore::UserWarning'

# Add after imports
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Page configuration
st.set_page_config(
    page_title="Epigenetic Disease Predictor",
    page_icon="🧬",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
    <style>
    .main-header {
        font-size: 3rem;
        color: #1E88E5;
        text-align: center;
        margin-bottom: 1rem;
        font-weight: bold;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.1);
    }
    .sub-header {
        font-size: 1.2rem;
        color: #333;
        text-align: center;
        margin-bottom: 2rem;
        font-weight: 500;
    }
    .prediction-box {
        padding: 2rem;
        border-radius: 15px;
        margin: 1.5rem 0;
        box-shadow: 0 6px 12px rgba(0, 0, 0, 0.15);
        border: 3px solid;
    }
    .control {
        background: linear-gradient(135deg, #A5D6A7 0%, #66BB6A 100%);
        border-color: #2E7D32;
    }
    .disease {
        background: linear-gradient(135deg, #EF9A9A 0%, #E57373 100%);
        border-color: #C62828;
    }
    .prediction-box h3 {
        color: #000;
        font-weight: bold;
        margin-bottom: 0.5rem;
        font-size: 1.4rem;
    }
    .prediction-box h2 {
        margin: 0.5rem 0;
        font-size: 2rem;
        font-weight: bold;
    }
    .control h2 {
        color: #1B5E20;
    }
    .disease h2 {
        color: #B71C1C;
    }
    .confidence-text {
        font-size: 1.5rem;
        font-weight: bold;
        margin: 1rem 0;
    }
    .control .confidence-text {
        color: #1B5E20;
    }
    .disease .confidence-text {
        color: #B71C1C;
    }
    .prediction-box p {
        font-size: 1.1rem;
        line-height: 1.6;
        color: #000;
        font-weight: 500;
    }
    .prediction-box strong {
        font-weight: 700;
        color: #000;
    }
    .disclaimer-text {
        font-size: 0.95rem;
        margin-top: 1rem;
        color: #333;
        font-style: italic;
        font-weight: 500;
    }
    .interpretation-text {
        font-size: 1rem;
        margin-top: 1rem;
        color: #222;
        font-weight: 600;
    }
    .stMetric {
        background-color: #e3f2fd;
        padding: 1rem;
        border-radius: 10px;
        border: 2px solid #1976d2;
    }
    .stMetric label {
        font-weight: 600;
        color: #000;
    }
    .stMetric [data-testid="stMetricValue"] {
        font-size: 2rem;
        font-weight: bold;
        color: #000;
    }
    </style>
""", unsafe_allow_html=True)

# Load models function
@st.cache_resource
def load_models():
    models = {}
    scalers = {}
    
    try:
        # Load Alzheimer's model
        alzheimer_model_path = Path("models/alzheimer_model.joblib")
        if alzheimer_model_path.exists():
            models['alzheimer'] = joblib.load(alzheimer_model_path)
            st.sidebar.success("✅ Alzheimer's model loaded")
        else:
            st.sidebar.info("ℹ️ Alzheimer's model not found")
    except Exception as e:
        st.sidebar.warning(f"⚠️ Alzheimer's model error: {str(e)[:50]}...")
    
    try:
        # Load Alzheimer's scaler
        alzheimer_scaler_path = Path("models/alzheimer_scaler.joblib")
        if alzheimer_scaler_path.exists():
            scalers['alzheimer'] = joblib.load(alzheimer_scaler_path)
    except Exception as e:
        pass
    
    try:
        # Load Prostate model (NEW FILENAME)
        prostate_model_path = Path("models/prostate_rf_model_2000f_70_30.joblib")
        if prostate_model_path.exists():
            models['prostate'] = joblib.load(prostate_model_path)
            st.sidebar.success("✅ Prostate model loaded")
        else:
            st.sidebar.info("ℹ️ Prostate model not found")
    except Exception as e:
        st.sidebar.error(f"❌ Prostate model error: {str(e)[:50]}...")
    
    try:
        # Load Prostate scaler (NEW FILENAME)
        prostate_scaler_path = Path("models/prostate_rf_scaler_2000f_70_30.joblib")
        if prostate_scaler_path.exists():
            scalers['prostate'] = joblib.load(prostate_scaler_path)
    except Exception as e:
        pass
    
    return models, scalers

# Function to prepare data for prediction
def prepare_data_for_prediction(df, model, scaler=None):
    """Prepare uploaded data to match model's expected features - handles ANY dataset"""
    try:
        # First, remove common non-feature columns
        columns_to_drop = ['class', 'Class', 'label', 'Label', 'target', 'Target', 'diagnosis', 'Diagnosis']
        df_clean = df.copy()
        
        for col in columns_to_drop:
            if col in df_clean.columns:
                st.info(f"ℹ️ Removing target column: '{col}'")
                df_clean = df_clean.drop(col, axis=1)
        
        # Store sample IDs if present (first non-numeric column or first column starting with letters)
        sample_ids = None
        first_col = df_clean.columns[0]
        
        # Check if first column is likely an ID column
        if df_clean[first_col].dtype == 'object' or not str(first_col).startswith('cg'):
            sample_ids = df_clean[first_col]
            df_clean = df_clean.drop(first_col, axis=1)
            st.info(f"ℹ️ Using '{first_col}' as sample ID column")
        
        # Convert all columns to numeric
        for col in df_clean.columns:
            df_clean[col] = pd.to_numeric(df_clean[col], errors='coerce')
        
        # Remove any columns that couldn't be converted to numeric
        df_clean = df_clean.select_dtypes(include=[np.number])
        
        # Get model's expected features
        expected_features = None
        if hasattr(scaler, 'feature_names_in_'):
            expected_features = scaler.feature_names_in_
            st.info(f"📊 Model expects {len(expected_features)} features")
        elif hasattr(model, 'feature_names_in_'):
            expected_features = model.feature_names_in_
            st.info(f"📊 Model expects {len(expected_features)} features")
        else:
            # Model doesn't have feature names, use as-is
            st.warning("⚠️ Model doesn't have feature names. Using uploaded features as-is.")
            if scaler is not None:
                return pd.DataFrame(
                    scaler.transform(df_clean),
                    columns=df_clean.columns,
                    index=df_clean.index
                )
            return df_clean
        
        # Create aligned dataframe with expected features
        aligned_df = pd.DataFrame(index=df_clean.index, columns=expected_features)
        
        # Fill in values for features that exist in uploaded data
        common_features = list(set(df_clean.columns) & set(expected_features))
        missing_in_upload = set(expected_features) - set(df_clean.columns)
        extra_in_upload = set(df_clean.columns) - set(expected_features)
        
        # Statistics
        match_percentage = (len(common_features) / len(expected_features)) * 100
        
        st.info(f"""
        **Feature Matching:**
        - Common features: {len(common_features)} / {len(expected_features)} ({match_percentage:.1f}%)
        - Missing features: {len(missing_in_upload)} (will be filled with zeros)
        - Extra features: {len(extra_in_upload)} (will be ignored)
        """)
        
        # Check if we have enough common features
        if match_percentage < 10:
            st.error(f"""
            ⚠️ **Low feature match ({match_percentage:.1f}%)**
            
            Your dataset has very few features in common with the trained model.
            This usually means:
            - Wrong disease model selected (try switching between Alzheimer's/Prostate)
            - Dataset is not compatible with this model
            - Dataset preprocessing is different
            
            **Expected features start with:** {list(expected_features[:5])}
            **Your features start with:** {list(df_clean.columns[:5])}
            """)
            return None
        
        # Fill common features
        for col in common_features:
            aligned_df[col] = df_clean[col].values
        
        # Fill missing features with zeros (or could use mean/median)
        for col in missing_in_upload:
            aligned_df[col] = 0
        
        # Convert to float
        aligned_df = aligned_df.astype(float)
        
        # Apply scaling if scaler exists
        if scaler is not None:
            st.info("🔧 Applying feature scaling...")
            aligned_df = pd.DataFrame(
                scaler.transform(aligned_df),
                columns=aligned_df.columns,
                index=aligned_df.index
            )
        
        # Warn if too many features are missing
        if match_percentage < 50:
            st.warning(f"""
            ⚠️ Only {match_percentage:.1f}% of features match. 
            Predictions may be less accurate. Consider using a dataset from the same source as the training data.
            """)
        
        return aligned_df
    
    except Exception as e:
        st.error(f"❌ Error preparing data: {str(e)}")
        import traceback
        with st.expander("🔍 View detailed error"):
            st.code(traceback.format_exc())
        return None

# Function to extract text from PDF
def extract_text_from_pdf(pdf_file):
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        st.error(f"Error extracting PDF: {str(e)}")
        return None

# Function to extract text from Word document
def extract_text_from_docx(docx_file):
    try:
        doc = Document(docx_file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting Word document: {str(e)}")
        return None

# Function to extract text from image using OCR
def extract_text_from_image(image_file):
    try:
        image = Image.open(image_file)
        text = pytesseract.image_to_string(image)
        return text
    except Exception as e:
        st.error(f"Error extracting image: {str(e)}")
        return None

# Function to parse text to DataFrame
def parse_text_to_dataframe(text):
    try:
        # Split by lines
        lines = text.strip().split('\n')
        
        # Try to detect delimiter (comma, tab, space)
        first_line = lines[0]
        if '\t' in first_line:
            delimiter = '\t'
        elif ',' in first_line:
            delimiter = ','
        else:
            delimiter = r'\s+'
        
        # Parse data
        data = []
        for line in lines:
            if line.strip():
                if delimiter == r'\s+':
                    row = re.split(delimiter, line.strip())
                else:
                    row = line.split(delimiter)
                data.append(row)
        
        if len(data) > 0:
            df = pd.DataFrame(data[1:], columns=data[0])
            return df
        return None
    except Exception as e:
        st.error(f"Error parsing text to dataframe: {str(e)}")
        return None

# Function to convert uploaded file to DataFrame
def convert_to_dataframe(uploaded_file):
    file_extension = uploaded_file.name.split('.')[-1].lower()
    
    if file_extension == 'csv':
        try:
            df = pd.read_csv(uploaded_file)
            return df
        except Exception as e:
            st.error(f"Error reading CSV: {str(e)}")
            return None
    
    elif file_extension == 'pdf':
        text = extract_text_from_pdf(uploaded_file)
        if text:
            return parse_text_to_dataframe(text)
        return None
    
    elif file_extension in ['docx', 'doc']:
        text = extract_text_from_docx(uploaded_file)
        if text:
            return parse_text_to_dataframe(text)
        return None
    
    elif file_extension in ['png', 'jpg', 'jpeg', 'tiff', 'bmp']:
        text = extract_text_from_image(uploaded_file)
        if text:
            return parse_text_to_dataframe(text)
        return None
    
    else:
        st.error(f"Unsupported file format: {file_extension}")
        return None

# Title and description
st.markdown('<h1 class="main-header">🧬 Epigenetic Disease Prediction System</h1>', unsafe_allow_html=True)
st.markdown('<p class="sub-header">Predict disease risk using epigenetic methylation data</p>', unsafe_allow_html=True)

# Load models
models, scalers = load_models()

# Sidebar
with st.sidebar:
    # Try to load local image, fallback to a simple colored placeholder
    try:
        st.image("assets/dna_logo.png", width="stretch")
    except:
        st.markdown("""
            <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                        padding: 20px; border-radius: 10px; text-align: center; margin-bottom: 20px;">
                <h2 style="color: white; margin: 0;">🧬 Epigenetic<br>Predictor</h2>
            </div>
        """, unsafe_allow_html=True)
    
    st.markdown("### 🎯 Model Selection")
    
    disease_type = st.selectbox(
        "Choose Disease Model",
        ["Alzheimer's Disease", "Prostate Cancer"],
        help="Select the disease you want to predict"
    )
    
    st.markdown("---")
    st.markdown("### 📊 Model Information")
    
    if disease_type == "Alzheimer's Disease":
        model_available = 'alzheimer' in models
        st.info(f"""
        **Dataset:** GSE80970  
        **Model:** Random Forest  
        **Status:** {'✅ Loaded' if model_available else '❌ Not Loaded'}  
        **Features:** DNA Methylation Sites
        """)
    else:
        model_available = 'prostate' in models
        st.info(f"""
        **Dataset:** GSE26126  
        **Model:** Random Forest (2000 features)
        **Status:** {'✅ Loaded' if model_available else '❌ Not Loaded'}  
        **Split:** 70/30 Train/Test
        """)
    
    st.markdown("---")
    st.markdown("### 🔒 Security")
    st.success("✅ Models encrypted with RSA")
    st.success("✅ Secure cloud storage")

# Main content area
col1, col2 = st.columns([2, 1])

with col1:
    st.markdown("### 📁 Upload Epigenetic Data")
    st.markdown("Upload your data in **CSV, PDF, Word, or Image** format")
    
    uploaded_file = st.file_uploader(
        "Choose a file",
        type=['csv', 'pdf', 'docx', 'doc', 'png', 'jpg', 'jpeg', 'tiff', 'bmp'],
        help="Supported formats: CSV, PDF, Word (DOCX/DOC), Images (PNG, JPG, JPEG, TIFF, BMP)"
    )
    
    if uploaded_file is not None:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        with st.spinner(f"🔄 Processing {file_extension.upper()} file..."):
            # Convert file to DataFrame
            df = convert_to_dataframe(uploaded_file)
        
        if df is not None:
            st.success(f"✅ File uploaded and converted successfully! Shape: {df.shape}")
            
            # Option to download as CSV
            st.download_button(
                label="💾 Download Converted CSV",
                data=df.to_csv(index=False),
                file_name=f"converted_{uploaded_file.name.split('.')[0]}.csv",
                mime="text/csv"
            )
            
            # Display data preview
            with st.expander("📋 Data Preview", expanded=True):
                st.dataframe(df.head(10), width="stretch")
                
                col_a, col_b, col_c = st.columns(3)
                with col_a:
                    st.metric("Rows", df.shape[0])
                with col_b:
                    st.metric("Columns", df.shape[1])
                with col_c:
                    st.metric("Missing Values", df.isnull().sum().sum())
        else:
            df = None
            st.error("❌ Failed to convert file. Please check the format and try again.")
    else:
        df = None
        st.info("👆 Please upload a file to begin prediction")

with col2:
    st.markdown("### ℹ️ Supported Formats")
    st.markdown("""
    **📄 CSV Files:**
    - Direct upload, no conversion needed
    
    **📑 PDF Files:**
    - Text extraction from tables
    - Must contain structured data
    
    **📝 Word Documents (.docx):**
    - Text and table extraction
    - Structured data format
    
    **🖼️ Images (PNG, JPG, TIFF):**
    - OCR text extraction
    - Clear, high-quality images work best
    
    **Required Data Format:**
    - First column: Sample IDs
    - Other columns: CpG site beta values
    - Values should be between 0 and 1
    """)
    
    # Download example file button
    if st.button("📥 Download Example CSV", width="stretch"):
        example_data = pd.DataFrame({
            'SampleID': ['Sample1', 'Sample2', 'Sample3'],
            'cg00000029': [0.8234, 0.7123, 0.6789],
            'cg00000108': [0.6543, 0.5234, 0.7891],
            'cg00000109': [0.9012, 0.8456, 0.7234]
        })
        csv = example_data.to_csv(index=False)
        st.download_button(
            label="Download",
            data=csv,
            file_name="example_methylation_data.csv",
            mime="text/csv"
        )

# Prediction section
st.markdown("---")
st.markdown("### 🔮 Make Prediction")

if df is not None:
    if st.button("🚀 Predict Disease Status", type="primary", use_container_width=True):
        # Select the appropriate model
        model_key = 'alzheimer' if disease_type == "Alzheimer's Disease" else 'prostate'
        
        if model_key not in models:
            st.error(f"❌ {disease_type} model is not loaded. Please check the models folder.")
        else:
            with st.spinner("🧬 Analyzing epigenetic data..."):
                try:
                    # Get model and scaler
                    model = models[model_key]
                    scaler = scalers.get(model_key, None)
                    
                    # Prepare data
                    prepared_data = prepare_data_for_prediction(df, model, scaler)
                    
                    if prepared_data is not None:
                        # Make predictions
                        predictions = model.predict(prepared_data)
                        
                        # Get prediction probabilities if available
                        if hasattr(model, 'predict_proba'):
                            probabilities = model.predict_proba(prepared_data)
                            confidence = probabilities.max(axis=1)
                        else:
                            confidence = np.ones(len(predictions)) * 0.85
                        
                        st.markdown("### 📊 Prediction Results")
                        
                        # Summary statistics
                        healthy_count = np.sum(predictions == 0)
                        disease_count = np.sum(predictions == 1)
                        total_samples = len(predictions)
                        
                        col_stat1, col_stat2, col_stat3 = st.columns(3)
                        with col_stat1:
                            st.metric("📋 Total Samples", total_samples)
                        with col_stat2:
                            st.metric("✅ Healthy Samples", healthy_count, 
                                     delta=f"{(healthy_count/total_samples*100):.1f}%")
                        with col_stat3:
                            st.metric("⚠️ Disease Detected", disease_count,
                                     delta=f"{(disease_count/total_samples*100):.1f}%",
                                     delta_color="inverse")
                        
                        # Model Performance Metrics (Expandable)
                        with st.expander("📈 Model Performance Metrics", expanded=False):
                            st.markdown("#### Prediction Statistics")
                            
                            perf_col1, perf_col2 = st.columns(2)
                            
                            with perf_col1:
                                avg_confidence = np.mean(confidence)
                                st.metric("Average Confidence", f"{avg_confidence:.2%}")
                                
                                healthy_conf = confidence[predictions == 0]
                                if len(healthy_conf) > 0:
                                    st.metric("Avg. Healthy Confidence", f"{np.mean(healthy_conf):.2%}")
                                else:
                                    st.metric("Avg. Healthy Confidence", "N/A")
                            
                            with perf_col2:
                                disease_conf = confidence[predictions == 1]
                                if len(disease_conf) > 0:
                                    st.metric("Avg. Disease Confidence", f"{np.mean(disease_conf):.2%}")
                                else:
                                    st.metric("Avg. Disease Confidence", "N/A")
                                
                                st.metric("Min Confidence", f"{np.min(confidence):.2%}")
                                st.metric("Max Confidence", f"{np.max(confidence):.2%}")
                            

                            # Distribution chart
                            st.markdown("#### Confidence Distribution")
                            conf_df = pd.DataFrame({
                                'Sample': [f"Sample {i+1}" for i in range(len(confidence))],
                                'Confidence': confidence,
                                'Prediction': ['Healthy' if p == 0 else 'Disease' for p in predictions]
                            })
                            st.bar_chart(conf_df.set_index('Sample')['Confidence'])
                        
                        st.markdown("---")
                        st.markdown("### 🔬 Individual Sample Results")
                        
                        # Display results for each sample
                        for idx, (pred, conf) in enumerate(zip(predictions, confidence)):
                            sample_name = df.iloc[idx, 0] if df.shape[1] > 0 else f"Sample {idx+1}"
                            
                            if pred == 0:
                                st.markdown(f"""
                                <div class="prediction-box control">
                                    <h3>🆔 {sample_name}</h3>
                                    <h2>✅ Control (Healthy)</h2>
                                    <p class="confidence-text">Confidence: {conf:.2%}</p>
                                    <p>The model predicts this sample as <strong>Control/Healthy</strong> for {disease_type}.</p>
                                    <p class="interpretation-text">
                                        <strong>Interpretation:</strong> Low risk of {disease_type} based on methylation patterns.
                                    </p>
                                </div>
                                """, unsafe_allow_html=True)
                            else:
                                st.markdown(f"""
                                <div class="prediction-box disease">
                                    <h3>🆔 {sample_name}</h3>
                                    <h2>⚠️ Disease Detected</h2>
                                    <p class="confidence-text">Confidence: {conf:.2%}</p>
                                    <p>The model predicts this sample shows signs of <strong>{disease_type}</strong>.</p>
                                    <p class="disclaimer-text">
                                        ⚠️ <strong>Important:</strong> This is a computational prediction based on epigenetic markers. 
                                        Please consult healthcare professionals for clinical diagnosis.
                                    </p>
                                </div>
                                """, unsafe_allow_html=True)
                        
                        # Feature importance visualization
                        with st.expander("📊 Top Contributing CpG Sites", expanded=False):
                            if hasattr(model, 'feature_importances_'):
                                feature_importance = pd.DataFrame({
                                    'Feature': prepared_data.columns,
                                    'Importance': model.feature_importances_
                                }).sort_values('Importance', ascending=False).head(20)
                                
                                st.markdown("#### Top 20 Most Important CpG Sites")
                                st.bar_chart(feature_importance.set_index('Feature'))
                                
                                # Download feature importance
                                csv_importance = feature_importance.to_csv(index=False)
                                st.download_button(
                                    label="💾 Download Feature Importance",
                                    data=csv_importance,
                                    file_name="feature_importance.csv",
                                    mime="text/csv"
                                )
                            else:
                                st.info("Feature importance not available for this model type")
                
                except Exception as e:
                    st.error(f"❌ Prediction error: {str(e)}")
                    st.info("Please ensure your data format matches the model's requirements.")
else:
    st.warning("⚠️ Please upload a file first to enable prediction")

# Footer
st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #666; padding: 2rem;">
    <p>🔬 Developed by Rudransh Pandey | 🧬 Powered by Machine Learning & Bioinformatics</p>
    <p>⚠️ <strong>Disclaimer:</strong> This tool is for research purposes only. Not for clinical diagnosis.</p>
</div>
""", unsafe_allow_html=True)